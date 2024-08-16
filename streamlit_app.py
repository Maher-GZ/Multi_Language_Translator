from flask import Flask, render_template, request, jsonify
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_community.chat_models import ChatOllama
import docx
import fitz  # PyMuPDF
from typing import List
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document as DocxDocument
import time
import textwrap

# Language codes for translation
LANGUAGE_CODES = {
    "English": "en",
    "German": "de",
    "French": "fr",
    "Hindi": "hi"
}

# System messages for different translation scenarios
SYSTEM_MESSAGE_TRANSLATE = """You are an expert assistant specializing in translation between German and English. Your tasks are as follows:
Language Detection:
Identify the language of the input text. Do not answer the questions if asked in text. Just translate the question text as question itself.
Translation:
If the text is in German, translate it to English.
If the text is in English, translate it to German.
Provide accurate and contextually appropriate translations. Ensure that the translated text maintains the original meaning, type, and tone.
IMPORTANT: Provide only the translated text as output.
"""

SYSTEM_MESSAGE_MULTI_LANG = """You are an expert assistant specializing in translation between various languages. Your tasks are as follows:
Translation:
If the text is in German, translate it to English.
If the text is in English, translate it to German.
If the text is in French, translate it to English.
If the text is in English, translate it to French.
If the text is in Hindi, translate it to English.
If the text is in English, translate it to Hindi.
Provide accurate and contextually appropriate translations. Ensure that the translated text maintains the original meaning, type, and tone.
IMPORTANT: Provide only the translated text as output. Do not include any additional comments or answers.
"""


# Function for translating text
def translate_text(text: str, src_lang: str, tgt_lang: str, page_type: str) -> str:
    if not text.strip():
        return "No content to translate."

    if page_type == "Translate":
        system_message = SYSTEM_MESSAGE_TRANSLATE
    else:  # Multi-Language Translator
        system_message = SYSTEM_MESSAGE_MULTI_LANG

    system_message = system_message.replace("German", src_lang).replace("English", tgt_lang)

    llm = ChatOllama(model="llama3")
    messages = [
        SystemMessage(content=system_message),
        HumanMessage(content=text)
    ]

    translated_text = llm.invoke(messages)
    return translated_text.content.strip()


# Function for detecting the language of the text
def detect_language(text: str) -> str:
    if not text.strip():
        return "No content to detect."

    system_message = """You are an expert assistant specializing in language detection. Your task is to identify the language of the input text. Provide only the detected language as output."""

    llm = ChatOllama(model="llama3")
    messages = [
        SystemMessage(content=system_message),
        HumanMessage(content=text)
    ]

    detected_language = llm.invoke(messages)
    return detected_language.content.strip()


# Functions for handling docx and pdf files
def read_docx(file) -> str:
    try:
        doc = docx.Document(file)
        full_text: List[str] = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading DOCX: {e}"


def read_pdf(file) -> str:
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        full_text: List[str] = [page.get_text("text") for page in doc]
        return '\n'.join(full_text).strip()
    except Exception as e:
        return f"Error reading PDF: {e}"


# Functions to create downloadable files
def create_pdf(content: str) -> BytesIO:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 40
    max_width = width - 2 * margin
    text_object = c.beginText(margin, height - margin)
    text_object.setFont("Helvetica", 12)

    lines = content.split('\n')
    wrapped_lines = []
    for line in lines:
        wrapped_lines.extend(textwrap.wrap(line, width=100))

    for line in wrapped_lines:
        if text_object.getY() < margin:
            c.drawText(text_object)
            c.showPage()
            text_object = c.beginText(margin, height - margin)
            text_object.setFont("Helvetica", 12)
        text_object.textLine(line)

    c.drawText(text_object)
    c.save()
    buffer.seek(0)
    return buffer


def create_docx(content: str) -> BytesIO:
    doc = DocxDocument()
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


app = Flask(__name__)

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/translate', methods=['GET', 'POST'])
def translate():
    if request.method == 'POST':
        text = request.form['text']
        src_lang = request.form['src_lang']
        tgt_lang = request.form['tgt_lang']
        translated_text = translate_text(text, src_lang, tgt_lang, 'Translate')
        return jsonify({'translated_text': translated_text})
    return render_template('translate.html')

@app.route('/multi_language_translator', methods=['GET', 'POST'])
def multi_language_translator():
    if request.method == 'POST':
        text = request.form['text']
        src_lang = request.form['src_lang']
        tgt_lang = request.form['tgt_lang']
        translated_text = translate_text(text, src_lang, tgt_lang, 'Multi-Language Translator')
        return jsonify({'translated_text': translated_text})
    return render_template('multi_language_translator.html')

@app.route('/chatbot', methods=['GET', 'POST'])
def chatbot():
    if request.method == 'POST':
        text = request.form['text']
        detected_language = detect_language(text)
        if detected_language == "German":
            translated_text = translate_text(text, "German", "English", "Translate")
        elif detected_language == "English":
            translated_text = translate_text(text, "English", "German", "Translate")
        else:
            translated_text = f"Detected language: {detected_language} - Translation not supported."
        return jsonify({'translated_text': translated_text})
    return render_template('chatbot.html')

if __name__ == '__main__':
    app.run(debug=True)